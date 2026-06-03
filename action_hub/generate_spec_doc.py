"""
Generate a compiled Word document from all ActionHub specification files.
Usage: python action_hub/generate_spec_doc.py
Output: specs/ActionHub_Specifications.docx
"""

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

PROJECT_ROOT = Path(__file__).resolve().parent.parent
SPECS_DIR = PROJECT_ROOT / "specs"
REQ_DIR = SPECS_DIR / "requirements"
OUTPUT_FILE = SPECS_DIR / "ActionHub_Specifications.docx"

FILES = [
    (SPECS_DIR / "context.md", "Project Context"),
    (REQ_DIR / "R00_initial_vision.md", None),
    (REQ_DIR / "R01_entities.md", None),
    (REQ_DIR / "R02_action_lifecycle.md", None),
    (REQ_DIR / "R03_assignment_workflow.md", None),
    (REQ_DIR / "R04_notifications.md", None),
    (REQ_DIR / "R05_dashboards_reporting.md", None),
    (REQ_DIR / "R06_security.md", None),
    (REQ_DIR / "R07_data_import.md", None),
    (REQ_DIR / "R08_taxonomy.md", None),
    (REQ_DIR / "R09_ui_content.md", None),
    (REQ_DIR / "R10_workflow_engine.md", None),
    (REQ_DIR / "R11_agent_framework.md", None),
    (REQ_DIR / "DECISIONS.md", None),
]


def setup_styles(doc):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    for index in range(1, 5):
        heading = doc.styles[f"Heading {index}"]
        heading.font.name = "Calibri"
        heading.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)

    doc.styles["Heading 1"].font.size = Pt(20)
    doc.styles["Heading 2"].font.size = Pt(16)
    doc.styles["Heading 3"].font.size = Pt(13)
    doc.styles["Heading 4"].font.size = Pt(11)


def add_cover_page(doc):
    for _ in range(6):
        doc.add_paragraph()

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("ActionHub")
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
    run.bold = True

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Centralized Action Log & Follow-Up Platform")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_paragraph()

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Requirements Specification - Compiled Document")
    run.font.size = Pt(14)
    run.bold = True

    doc.add_paragraph()

    metadata = [
        ("Organization", "ActionHub Organization"),
        ("Version", "V1 MVP (1.5-day sprint)"),
        ("Date", "February 2026"),
        ("Status", "Requirements-level specification"),
        ("Classification", "Internal - For stakeholder review"),
    ]
    table = doc.add_table(rows=len(metadata), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row_index, (key, value) in enumerate(metadata):
        row = table.rows[row_index]
        row.cells[0].text = key
        row.cells[1].text = value
        if row.cells[0].paragraphs[0].runs:
            row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_page_break()


def add_toc_page(doc):
    doc.add_heading("Table of Contents", level=1)
    entries = [
        "1. Project Context",
        "2. R00 - Vision & Foundation",
        "3. R01 - Entity Inventory",
        "4. R02 - Action Lifecycle & Business Logic",
        "5. R03 - Assignment & RACI Workflow",
        "6. R04 - Notifications & Follow-Up",
        "7. R05 - Dashboards & Reporting",
        "8. R06 - Security & Access Control",
        "9. R07 - Data Import (Excel Seed)",
        "10. R08 - Taxonomy & Classification",
        "11. R09 - UI/UX & Bilingual Interface",
        "12. R10 - Workflow Engine (V2+ Vision)",
        "13. R11 - Agent Framework (V3+ Vision)",
        "14. Decision Log (D1-D165)",
    ]
    for entry in entries:
        paragraph = doc.add_paragraph(entry)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.space_before = Pt(0)

    doc.add_page_break()


def parse_table(lines, start_index):
    rows = []
    index = start_index
    while index < len(lines) and "|" in lines[index]:
        line = lines[index].strip()
        if line.startswith("|"):
            cells = [cell.strip() for cell in line.split("|")[1:-1]]
            if cells and not all(re.match(r"^[-:]+$", cell) for cell in cells):
                rows.append(cells)
        index += 1
    return rows, index


def add_table_to_doc(doc, rows):
    if not rows:
        return

    num_cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols, style="Light Grid Accent 1")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for row_index, row_data in enumerate(rows):
        row = table.rows[row_index]
        for col_index, cell_text in enumerate(row_data):
            if col_index >= num_cols:
                continue
            cell = row.cells[col_index]
            cell.text = ""
            paragraph = cell.paragraphs[0]
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", cell_text.strip())
            text = re.sub(r"`(.+?)`", r"\1", text)
            run = paragraph.add_run(text)
            run.font.size = Pt(9)
            if row_index == 0:
                run.bold = True

    doc.add_paragraph()


def process_markdown_file(doc, filepath, override_title=None):
    with open(filepath, "r", encoding="utf-8") as file_handle:
        lines = file_handle.read().split("\n")

    index = 0
    first_h1_seen = False
    in_code_block = False
    code_lines = []

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code_block:
                code_text = "\n".join(code_lines)
                if code_text.strip():
                    paragraph = doc.add_paragraph()
                    paragraph.paragraph_format.left_indent = Cm(0.5)
                    paragraph.paragraph_format.space_before = Pt(4)
                    paragraph.paragraph_format.space_after = Pt(4)
                    run = paragraph.add_run(code_text)
                    run.font.name = "Consolas"
                    run.font.size = Pt(8)
                    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                    shading = run._element.get_or_add_rPr()
                    shading_element = shading.makeelement(qn("w:shd"), {
                        qn("w:val"): "clear",
                        qn("w:color"): "auto",
                        qn("w:fill"): "F3F4F6",
                    })
                    shading.append(shading_element)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            index += 1
            continue

        if in_code_block:
            code_lines.append(line)
            index += 1
            continue

        if not stripped or stripped in ("---", "***", "___"):
            index += 1
            continue

        heading_match = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            text = re.sub(r"`(.+?)`", r"\1", text)
            if level == 1 and not first_h1_seen:
                first_h1_seen = True
                doc.add_heading(override_title or text, level=1)
            else:
                doc.add_heading(text, level=min(level, 4))
            index += 1
            continue

        if stripped.startswith(">"):
            quote_text = stripped.lstrip("> ").strip()
            while index + 1 < len(lines) and lines[index + 1].strip().startswith(">"):
                index += 1
                quote_text += " " + lines[index].strip().lstrip("> ").strip()
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(1)
            quote_text = re.sub(r"\*\*(.+?)\*\*", r"\1", quote_text)
            quote_text = re.sub(r"`(.+?)`", r"\1", quote_text)
            run = paragraph.add_run(quote_text)
            run.font.size = Pt(9)
            run.italic = True
            run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
            index += 1
            continue

        if stripped.startswith("|") and "|" in stripped:
            rows, new_index = parse_table(lines, index)
            add_table_to_doc(doc, rows)
            index = new_index
            continue

        bullet_match = re.match(r"^(\s*)[-*+]\s+(.+)$", stripped)
        if bullet_match:
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", bullet_match.group(2))
            text = re.sub(r"`(.+?)`", r"\1", text)
            paragraph = doc.add_paragraph(text, style="List Bullet")
            paragraph.paragraph_format.space_after = Pt(2)
            if paragraph.runs:
                paragraph.runs[0].font.size = Pt(10)
            index += 1
            continue

        num_match = re.match(r"^(\s*)\d+[.)]\s+(.+)$", stripped)
        if num_match:
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", num_match.group(2))
            text = re.sub(r"`(.+?)`", r"\1", text)
            paragraph = doc.add_paragraph(text, style="List Number")
            paragraph.paragraph_format.space_after = Pt(2)
            index += 1
            continue

        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(4)
        for segment in re.split(r"(\*\*.+?\*\*)", stripped):
            bold_match = re.match(r"\*\*(.+?)\*\*", segment)
            if bold_match:
                run = paragraph.add_run(bold_match.group(1))
                run.bold = True
            else:
                run = paragraph.add_run(re.sub(r"`(.+?)`", r"\1", segment))
            run.font.size = Pt(10)

        index += 1


def add_section_number(doc, number, total):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(f"- Section {number} of {total} -")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def main():
    print("Generating ActionHub Specifications Document...")

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    setup_styles(doc)
    add_cover_page(doc)
    add_toc_page(doc)

    total = len(FILES)
    for index, (filepath, override_title) in enumerate(FILES, 1):
        if not filepath.exists():
            print(f"  WARNING: {filepath} not found, skipping")
            continue
        print(f"  [{index}/{total}] Processing {filepath.name}...")
        add_section_number(doc, index, total)
        process_markdown_file(doc, filepath, override_title)
        if index < total:
            doc.add_page_break()

    doc.save(str(OUTPUT_FILE))
    print(f"\nDone! Saved to: {OUTPUT_FILE}")
    print(f"File size: {OUTPUT_FILE.stat().st_size / 1024:.0f} KB")


if __name__ == "__main__":
    main()